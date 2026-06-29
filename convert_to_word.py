#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown测试文档转换为Word文档
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_chinese_font(run, font_name='宋体', size=12):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)


def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_chinese_font(run, '黑体', 18 - level * 2)


def add_paragraph(doc, text, bold=False, font_size=12):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, '宋体', font_size)
    run.bold = bold
    return para


def add_table(doc, data, headers=None):
    """添加表格"""
    if not data:
        return
    
    rows = len(data)
    cols = len(data[0]) if data else 0
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # 填充表头
    if headers:
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(header)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, '黑体', 11)
                    run.bold = True
    
    # 填充数据
    start_row = 0 if not headers else 1
    for i in range(start_row, rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = str(data[i][j])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, '宋体', 10)


def parse_markdown_to_word(md_file, docx_file):
    """解析Markdown并转换为Word"""
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_table = False
    table_data = []
    table_headers = []
    in_code_block = False
    code_content = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()
        
        # 代码块处理
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_content = []
            else:
                in_code_block = False
                if code_content:
                    para = doc.add_paragraph('代码示例：', style='Intense Quote')
                    for run in para.runs:
                        set_chinese_font(run, '黑体', 10)
                    code_para = doc.add_paragraph('\n'.join(code_content))
                    for run in code_para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # 表格处理
        if stripped.startswith('|') and '|' in stripped[1:]:
            if not in_table:
                in_table = True
                table_data = []
                table_headers = []
            
            # 解析表格行
            cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
            table_data.append(cells)
            
            # 检查是否是分隔线
            if i + 1 < len(lines):
                next_line = lines[i + 1].rstrip('\n').strip()
                if next_line.startswith('|') and all(c in '|-: ' for c in next_line.replace('|', '')):
                    table_headers = table_data[0]
                    table_data = []
                    i += 1  # 跳过分隔线
            
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                if table_data:
                    add_table(doc, table_data, table_headers)
        
        # 标题处理
        if stripped.startswith('#'):
            level = len(stripped.split()[0])
            text = stripped[level:].strip()
            add_heading(doc, text, min(level, 3))
            i += 1
            continue
        
        # 分隔线
        if stripped.startswith('---') or stripped.startswith('==='):
            doc.add_paragraph()
            i += 1
            continue
        
        # 普通段落
        if stripped:
            # 处理粗体
            if stripped.startswith('**') and stripped.endswith('**'):
                text = stripped[2:-2]
                add_paragraph(doc, text, bold=True)
            else:
                add_paragraph(doc, line)
        else:
            doc.add_paragraph()
        
        i += 1
    
    # 处理最后可能遗留的表格
    if in_table and table_data:
        add_table(doc, table_data, table_headers)
    
    # 保存文档
    doc.save(docx_file)
    print(f'Word文档已生成: {docx_file}')


if __name__ == '__main__':
    md_file = os.path.join(os.path.dirname(__file__), 'docs', 'software-testing-document.md')
    docx_file = os.path.join(os.path.dirname(__file__), '校园二手物品交易系统_软件测试文档.docx')
    
    parse_markdown_to_word(md_file, docx_file)
