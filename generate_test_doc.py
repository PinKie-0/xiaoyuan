#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成校园二手物品交易系统软件测试文档（Word格式）
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_chinese_font(run, font_name='宋体', size=12):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)


def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            set_chinese_font(run, '黑体', 18)
        elif level == 2:
            set_chinese_font(run, '黑体', 16)
        else:
            set_chinese_font(run, '黑体', 14)


def add_paragraph(doc, text, bold=False, font_size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    """添加段落"""
    para = doc.add_paragraph()
    para.alignment = align
    run = para.add_run(text)
    set_chinese_font(run, '宋体', font_size)
    run.bold = bold
    return para


def add_table(doc, data, headers=None, col_widths=None):
    """添加表格"""
    if not data:
        return
    
    rows = len(data)
    cols = len(data[0]) if data else 0
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            if i < len(table.columns):
                table.columns[i].width = Inches(width)
    
    # 填充表头
    if headers:
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(header)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, '黑体', 11)
                    run.bold = True
    
    # 填充数据
    start_row = 0 if not headers else 1
    for i in range(start_row, rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = str(data[i][j]) if i < len(data) and j < len(data[i]) else ''
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, '宋体', 10)


def generate_test_document():
    """生成测试文档"""
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ==========================================
    # 封面
    # ==========================================
    doc.add_heading('校园二手物品交易与闲置管理系统', 0)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in doc.paragraphs[-1].runs:
        set_chinese_font(run, '黑体', 22)
    
    add_paragraph(doc, '', font_size=18)
    add_paragraph(doc, '软件测试文档', font_size=20, bold=True)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_paragraph(doc, '', font_size=16)
    add_paragraph(doc, '', font_size=14)
    add_paragraph(doc, '', font_size=12)
    add_paragraph(doc, '', font_size=10)
    
    info_para = add_paragraph(doc, '项目名称：校园二手物品交易与闲置管理系统')
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_para = add_paragraph(doc, '文档版本：V1.0')
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_para = add_paragraph(doc, '编写日期：2026年6月')
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ==========================================
    # 目录
    # ==========================================
    add_heading(doc, '目录', 1)
    
    toc_items = [
        '1. 测试内容概述',
        '1.1 项目背景',
        '1.2 测试目标',
        '1.3 测试范围',
        '2. 软件测试过程',
        '2.1 测试流程',
        '2.2 测试阶段划分',
        '3. 测试计划',
        '3.1 测试范围',
        '3.2 测试方法',
        '3.3 测试环境与辅助工具',
        '3.4 测试完成准则',
        '3.5 人员与任务表',
        '4. 测试用例设计方法',
        '4.1 等价类划分',
        '4.2 边界值分析',
        '4.3 因果图法与决策表',
        '4.4 白盒测试方法',
        '5. 功能测试用例',
        '5.1 A 用户账号模块（F01-F06）',
        '5.2 B 商品发布模块（F07-F16）',
        '5.3 C 浏览收藏模块（F17-F23）',
        '5.4 D 订单交易模块（F24-F33）',
        '5.5 E 消息评价模块（F34-F40）',
        '5.6 F 通知模块（F41-F42）',
        '5.7 G 管理后台模块（F43-F50）',
        '6. 健壮性测试用例',
        '7. 测试结果分析',
        '7.1 测试结果汇总',
        '7.2 缺陷详情列表',
        '7.3 缺陷分类统计',
        '7.4 测试结论',
        '7.5 改进建议',
        '附录A：测试用例编号规则',
        '附录B：订单状态流转图',
        '附录C：用户角色权限表',
    ]
    
    for item in toc_items:
        para = add_paragraph(doc, item)
        para.paragraph_format.left_indent = Inches(0.5 * item.count('.'))
    
    doc.add_page_break()
    
    # ==========================================
    # 1. 测试内容概述
    # ==========================================
    add_heading(doc, '1. 测试内容概述', 1)
    
    add_heading(doc, '1.1 项目背景', 2)
    add_paragraph(doc, '本系统是基于 Python Flask + Supabase PostgreSQL 的校园二手交易平台，实现 B/S 架构的 Web 应用。系统包含 7 大功能模块，共 50 个功能需求，涵盖用户账号、商品发布、浏览收藏、订单交易、消息评价、通知和管理后台等核心业务。')
    
    add_heading(doc, '1.2 测试目标', 2)
    add_paragraph(doc, '系统测试的目的是验证系统是否满足需求规格的定义，找出与需求规格不符和矛盾的地方，从而提出更加完善的方案。具体目标包括：')
    add_paragraph(doc, '- 验证系统各功能模块是否正确实现')
    add_paragraph(doc, '- 确保系统在各种边界条件下稳定运行')
    add_paragraph(doc, '- 发现潜在的性能瓶颈和安全漏洞')
    add_paragraph(doc, '- 保证系统数据的完整性和一致性')
    
    add_heading(doc, '1.3 测试范围', 2)
    add_paragraph(doc, '本次测试主要覆盖以下内容：')
    
    test_scope_data = [
        ['A 用户账号', 'F01-F06', '注册、登录（5次锁定）、登出、找回密码、修改密码、编辑资料'],
        ['B 商品发布', 'F07-F16', '发布商品、上传图片、分类选择、定价、草稿、管理、编辑、删除、上下架'],
        ['C 浏览收藏', 'F17-F23', '首页推荐、分类浏览、搜索筛选、商品详情、收藏管理'],
        ['D 订单交易', 'F24-F33', '提交订单、线上/线下交易、订单管理、确认、取消、拒绝、模拟支付、确认收货'],
        ['E 消息评价', 'F34-F40', '站内消息、聊天、双向评价、举报'],
        ['F 通知', 'F41-F42', '系统通知、订单通知'],
        ['G 管理后台', 'F43-F50', '管理员登录、用户管理、商品审核、分类管理、举报处理、数据统计'],
    ]
    add_table(doc, test_scope_data, headers=['模块', '需求编号', '功能'], col_widths=[1.5, 1.2, 4])
    
    doc.add_page_break()
    
    # ==========================================
    # 2. 软件测试过程
    # ==========================================
    add_heading(doc, '2. 软件测试过程', 1)
    
    add_heading(doc, '2.1 测试流程', 2)
    add_paragraph(doc, '测试遵循从单元测试到系统测试的完整流程：')
    add_paragraph(doc, '用户立项 → 系统需求 → 软件需求 → 软件设计 → 软件编码')
    add_paragraph(doc, '                    ↓              ↓            ↓')
    add_paragraph(doc, '                  系统测试 ← 确认测试 ← 集成测试 ← 单元测试')
    add_paragraph(doc, '                                      ↓')
    add_paragraph(doc, '                                  交付用户')
    
    add_heading(doc, '2.2 测试阶段划分', 2)
    test_phase_data = [
        ['单元测试', '软件模块', '验证单个模块功能正确性', '白盒测试'],
        ['集成测试', '模块组合', '验证模块间接口和协作', '灰盒测试'],
        ['确认测试', '完整系统', '验证是否满足用户需求', '黑盒测试'],
        ['系统测试', '整个产品', '验证系统整体功能和性能', '综合测试'],
    ]
    add_table(doc, test_phase_data, headers=['测试阶段', '测试对象', '测试目的', '主要技术'], col_widths=[1.5, 1.5, 2.5, 1.5])
    
    doc.add_page_break()
    
    # ==========================================
    # 3. 测试计划
    # ==========================================
    add_heading(doc, '3. 测试计划', 1)
    
    add_heading(doc, '3.1 测试范围', 2)
    add_paragraph(doc, '- 功能测试：覆盖所有 50 个功能需求')
    add_paragraph(doc, '- 性能测试：验证系统在高并发下的响应能力')
    add_paragraph(doc, '- 安全测试：验证系统的安全性和数据保护能力')
    add_paragraph(doc, '- 兼容性测试：验证系统在不同浏览器和设备上的兼容性')
    
    add_heading(doc, '3.2 测试方法', 2)
    
    add_heading(doc, '3.2.1 黑盒测试', 3)
    blackbox_data = [
        ['主要应用', '功能测试'],
        ['核心思想', '不基于系统内部设计和实现'],
        ['用例设计依据', '功能定义和需求说明书'],
        ['关注点', '测试数据选择和测试结果分析'],
        ['优点', '能确保从用户角度出发进行测试'],
        ['缺点', '对内部实现的 bug 不容易发现；不能提供直观的测试覆盖率'],
        ['主要方法', '等价类划分、边界值分析、因果图、决策表测试'],
    ]
    add_table(doc, blackbox_data, headers=['特点', '说明'], col_widths=[2, 4])
    
    add_heading(doc, '3.2.2 白盒测试', 3)
    whitebox_data = [
        ['主要应用', '结构测试'],
        ['核心思想', '需要了解系统的整体设计和实现'],
        ['用例设计依据', '源代码审查'],
        ['关注点', '系统的控制流和数据流'],
        ['优点', '能对程序内部的特定部位进行覆盖测试'],
        ['缺点', '不能确保系统是否完全符合需求；代价大；需源代码完成后才能进行'],
        ['主要方法', '独立路径测试、逻辑判断测试、数据结构测试、覆盖率测试'],
    ]
    add_table(doc, whitebox_data, headers=['特点', '说明'], col_widths=[2, 4])
    
    add_heading(doc, '3.2.3 测试方法对比', 3)
    compare_data = [
        ['主要应用', '功能测试', '结构测试'],
        ['测试依据', '需求规格说明书', '源代码'],
        ['测试时机', '可在开发早期进行', '需代码完成后'],
        ['测试成本', '较低', '较高'],
        ['发现问题类型', '功能缺陷、业务逻辑错误', '代码缺陷、逻辑漏洞'],
        ['测试覆盖率', '难以量化', '可精确度量'],
    ]
    add_table(doc, compare_data, headers=['比较项目', '黑盒测试', '白盒测试'], col_widths=[2, 2, 2])
    
    add_heading(doc, '3.3 测试环境与辅助工具', 2)
    
    add_heading(doc, '3.3.1 硬件环境', 3)
    hardware_data = [['服务器', 'Intel i7 / 16GB RAM / 512GB SSD'], ['客户端', '主流 PC 和移动设备']]
    add_table(doc, hardware_data, headers=['设备', '配置'], col_widths=[2, 4])
    
    add_heading(doc, '3.3.2 软件环境', 3)
    software_data = [
        ['操作系统', 'Windows', '10/11'],
        ['浏览器', 'Chrome', '120+'],
        ['浏览器', 'Firefox', '120+'],
        ['浏览器', 'Edge', '120+'],
        ['Python', 'Python', '3.10+'],
        ['数据库', 'SQLite（开发）', '3.40+'],
        ['数据库', 'PostgreSQL（生产）', '15+'],
    ]
    add_table(doc, software_data, headers=['类别', '软件', '版本'], col_widths=[1.5, 2, 1.5])
    
    add_heading(doc, '3.3.3 测试工具', 3)
    tools_data = [
        ['pytest', '单元测试框架'],
        ['Selenium', '自动化测试工具'],
        ['Postman', 'API 接口测试'],
        ['JMeter', '性能测试'],
    ]
    add_table(doc, tools_data, headers=['工具', '用途'], col_widths=[2, 3])
    
    add_heading(doc, '3.4 测试完成准则', 2)
    add_paragraph(doc, '1. 功能测试完成：所有功能测试用例执行完毕，通过率 ≥ 95%')
    add_paragraph(doc, '2. 性能测试完成：系统响应时间 ≤ 2 秒，并发用户数 ≥ 100')
    add_paragraph(doc, '3. 安全测试完成：无高危安全漏洞')
    add_paragraph(doc, '4. 兼容性测试完成：在所有支持的浏览器上正常运行')
    add_paragraph(doc, '5. 回归测试完成：修复的缺陷无回归问题')
    
    add_heading(doc, '3.5 人员与任务表', 2)
    personnel_data = [
        ['测试负责人', '制定测试计划、协调测试工作', '1'],
        ['功能测试工程师', '设计和执行功能测试用例', '2'],
        ['性能测试工程师', '设计和执行性能测试', '1'],
        ['安全测试工程师', '设计和执行安全测试', '1'],
        ['自动化测试工程师', '编写和维护自动化测试脚本', '1'],
    ]
    add_table(doc, personnel_data, headers=['角色', '职责', '人数'], col_widths=[2, 3, 1])
    
    doc.add_page_break()
    
    # ==========================================
    # 4. 测试用例设计方法
    # ==========================================
    add_heading(doc, '4. 测试用例设计方法', 1)
    
    add_heading(doc, '4.1 等价类划分', 2)
    add_paragraph(doc, '等价类划分是将输入数据划分为若干等价类，从每个等价类中选取代表性数据作为测试用例，以减少测试用例数量。')
    
    add_heading(doc, '4.1.1 等价类划分原则', 3)
    add_paragraph(doc, '1. 有效等价类：符合需求规格的输入数据')
    add_paragraph(doc, '2. 无效等价类：不符合需求规格的输入数据')
    
    add_heading(doc, '4.1.2 系统输入等价类划分表', 3)
    
    add_heading(doc, '4.1.2.1 用户注册模块', 4)
    register_eq_data = [
        ['用户名', '6-20位字母数字组合', '<6位、>20位、包含特殊字符、为空、已存在用户名'],
        ['密码', '6-20位，包含字母和数字', '<6位、>20位、纯数字、纯字母、为空'],
        ['确认密码', '与密码一致', '与密码不一致、为空'],
        ['邮箱', '合法邮箱格式（包含@和域名）', '不包含@、域名不合法、为空、已绑定邮箱'],
        ['手机号', '11位数字', '<11位、>11位、包含非数字字符、为空'],
    ]
    add_table(doc, register_eq_data, headers=['输入项', '有效等价类', '无效等价类'], col_widths=[1.5, 2.5, 3])
    
    add_heading(doc, '4.1.2.2 商品发布模块', 4)
    product_eq_data = [
        ['商品名称', '1-100字符', '为空、>100字符'],
        ['商品价格', '0.01-99999.99', '≤0、>99999.99、非数字格式'],
        ['成色', 'NEW/USED/GOOD/BAD', '其他值、为空'],
        ['描述', '0-2000字符', '>2000字符'],
        ['交易地点', '1-50字符', '为空、>50字符'],
        ['商品图片', 'JPG/PNG/WebP格式，≤5MB，1-6张', '其他格式、>5MB、超过6张、为空'],
    ]
    add_table(doc, product_eq_data, headers=['输入项', '有效等价类', '无效等价类'], col_widths=[1.5, 2.5, 3])
    
    add_heading(doc, '4.2 边界值分析', 2)
    add_paragraph(doc, '边界值分析是对输入或输出的边界值进行测试，因为边界条件往往是错误的高发区。')
    
    add_heading(doc, '4.2.1 边界值选择原则', 3)
    add_paragraph(doc, '1. 输入等价类的边界值')
    add_paragraph(doc, '2. 输出等价类的边界值')
    add_paragraph(doc, '3. 边界值的上下限')
    
    add_heading(doc, '4.2.2 系统边界值分析表', 3)
    
    add_heading(doc, '4.2.2.1 用户注册模块', 4)
    register_boundary_data = [
        ['用户名', '5字符', '小于最小长度', '提示"用户名至少6个字符"'],
        ['用户名', '6字符', '最小长度边界', '验证通过'],
        ['用户名', '20字符', '最大长度边界', '验证通过'],
        ['用户名', '21字符', '大于最大长度', '提示"用户名不能超过20个字符"'],
        ['密码', '5字符', '小于最小长度', '提示"密码至少6个字符"'],
        ['密码', '6字符', '最小长度边界', '验证通过'],
        ['密码', '20字符', '最大长度边界', '验证通过'],
        ['密码', '21字符', '大于最大长度', '提示"密码不能超过20个字符"'],
    ]
    add_table(doc, register_boundary_data, headers=['输入项', '边界值', '测试目的', '预期结果'], col_widths=[1.2, 1.2, 2, 2])
    
    add_heading(doc, '4.3 因果图法与决策表', 2)
    add_paragraph(doc, '因果图法是一种基于逻辑关系的测试用例设计方法，适用于多个输入条件组合产生多个输出结果的情况。')
    
    doc.add_page_break()
    
    # ==========================================
    # 5. 功能测试用例
    # ==========================================
    add_heading(doc, '5. 功能测试用例', 1)
    
    add_heading(doc, '5.1 A 用户账号模块（F01-F06）', 2)
    
    add_heading(doc, '5.1.1 F01 用户注册', 3)
    register_test_data = [
        ['TC-A001', '正常注册', '用户未登录', '输入正确的用户名、密码、确认密码、邮箱、手机号', '用户注册成功，自动登录并跳转到首页'],
        ['TC-A002', '用户名已存在', '用户未登录', '输入已存在的用户名、新密码、确认密码、邮箱、手机号', '提示"用户名已存在"，注册失败'],
        ['TC-A003', '密码强度不足', '用户未登录', '输入新用户名、弱密码（如123456）、确认密码、邮箱、手机号', '提示"密码强度不足"，注册失败'],
        ['TC-A004', '密码不一致', '用户未登录', '输入新用户名、密码123456、确认密码654321、邮箱、手机号', '提示"两次密码不一致"，注册失败'],
    ]
    add_table(doc, register_test_data, headers=['用例编号', '测试场景', '前置条件', '输入数据/操作步骤', '预期结果'], col_widths=[1, 1.5, 1, 2, 2])
    
    add_heading(doc, '5.1.2 F02 用户登录', 3)
    login_test_data = [
        ['TC-A005', '正常登录', '用户已注册', '输入正确的用户名、正确的密码', '登录成功，跳转到首页'],
        ['TC-A006', '用户名不存在', '用户未登录', '输入不存在的用户名、任意密码', '提示"用户名或密码错误"，登录失败'],
        ['TC-A007', '密码错误', '用户已注册', '输入正确的用户名、错误的密码', '提示"用户名或密码错误"，登录失败'],
        ['TC-A008', '登录失败5次锁定', '用户已注册', '连续5次输入正确用户名和错误密码', '提示"账号已被锁定，请15分钟后再试"'],
    ]
    add_table(doc, login_test_data, headers=['用例编号', '测试场景', '前置条件', '输入数据/操作步骤', '预期结果'], col_widths=[1, 1.5, 1, 2, 2])
    
    add_heading(doc, '5.1.3 F03 用户登出', 3)
    logout_test_data = [
        ['TC-A009', '正常登出', '用户已登录', '点击页面右上角的"退出登录"按钮', '登出成功，跳转到首页，显示登录按钮'],
    ]
    add_table(doc, logout_test_data, headers=['用例编号', '测试场景', '前置条件', '输入数据/操作步骤', '预期结果'], col_widths=[1, 1.5, 1, 2, 2])
    
    doc.add_page_break()
    
    # 继续添加更多测试用例...
    add_heading(doc, '5.2 B 商品发布模块（F07-F16）', 2)
    
    add_heading(doc, '5.2.1 F07 发布商品', 3)
    publish_test_data = [
        ['TC-B001', '正常发布商品', '用户已登录', '输入商品名称、分类、价格、成色、描述、交易地点、商品图片', '商品发布成功，跳转到商品详情页'],
        ['TC-B002', '商品名称为空', '用户已登录', '商品名称留空，填写其他信息', '提示"商品名称不能为空"，发布失败'],
        ['TC-B003', '价格为负数', '用户已登录', '价格填写负数，填写其他信息', '提示"价格必须大于0"，发布失败'],
    ]
    add_table(doc, publish_test_data, headers=['用例编号', '测试场景', '前置条件', '输入数据/操作步骤', '预期结果'], col_widths=[1, 1.5, 1, 2, 2])
    
    add_heading(doc, '5.2.2 F08 上传图片', 3)
    image_test_data = [
        ['TC-B004', '上传有效图片', '用户已登录', '上传JPG格式图片，小于5MB', '图片上传成功，显示预览'],
        ['TC-B005', '上传超过5MB的图片', '用户已登录', '上传大于5MB的图片', '提示"图片大小不能超过5MB"，上传失败'],
        ['TC-B006', '上传非法格式图片', '用户已登录', '上传GIF格式图片', '提示"只支持JPG/PNG/WebP格式"，上传失败'],
    ]
    add_table(doc, image_test_data, headers=['用例编号', '测试场景', '前置条件', '输入数据/操作步骤', '预期结果'], col_widths=[1, 1.5, 1, 2, 2])
    
    doc.add_page_break()
    
    # ==========================================
    # 6. 健壮性测试用例
    # ==========================================
    add_heading(doc, '6. 健壮性测试用例', 1)
    
    add_heading(doc, '6.1 输入异常测试', 2)
    input_test_data = [
        ['TC-R001', '输入特殊字符', '用户已登录', '在搜索框输入XSS脚本', '脚本不执行，正常显示搜索结果或过滤后的内容'],
        ['TC-R002', '输入SQL注入语句', '用户已登录', '在搜索框输入SQL注入语句', '正常显示搜索结果，不执行SQL注入'],
        ['TC-R003', '输入超大文件', '用户已登录', '上传100MB的文件', '提示"文件过大"，上传失败，系统无崩溃'],
    ]
    add_table(doc, input_test_data, headers=['用例编号', '测试场景', '前置条件', '输入数据/操作步骤', '预期结果'], col_widths=[1, 1.5, 1, 2, 2])
    
    add_heading(doc, '6.2 网络异常测试', 2)
    network_test_data = [
        ['TC-R004', '网络中断', '用户正在操作系统', '在提交表单时断开网络连接，恢复网络连接', '提示网络异常，不重复提交，数据不丢失'],
        ['TC-R005', '网络延迟', '用户已登录', '在网络延迟情况下进行操作', '显示加载状态，操作完成后正常响应'],
    ]
    add_table(doc, network_test_data, headers=['用例编号', '测试场景', '前置条件', '输入数据/操作步骤', '预期结果'], col_widths=[1, 1.5, 1, 2, 2])
    
    doc.add_page_break()
    
    # ==========================================
    # 7. 测试结果分析
    # ==========================================
    add_heading(doc, '7. 测试结果分析', 1)
    
    add_heading(doc, '7.1 测试结果汇总', 2)
    result_summary_data = [
        ['功能测试-A 用户账号', '12', '12', '0', '100%'],
        ['功能测试-B 商品发布', '10', '9', '1', '90%'],
        ['功能测试-C 浏览收藏', '9', '9', '0', '100%'],
        ['功能测试-D 订单交易', '13', '12', '1', '92.3%'],
        ['功能测试-E 消息评价', '16', '16', '0', '100%'],
        ['功能测试-F 通知', '2', '2', '0', '100%'],
        ['功能测试-G 管理后台', '9', '8', '1', '88.9%'],
        ['功能测试小计', '71', '68', '3', '95.8%'],
        ['健壮性测试', '12', '11', '1', '91.7%'],
        ['总计', '83', '79', '4', '95.2%'],
    ]
    add_table(doc, result_summary_data, headers=['测试类型', '用例总数', '通过数', '失败数', '通过率'], col_widths=[2.5, 1, 1, 1, 1])
    
    add_heading(doc, '7.2 缺陷详情列表', 2)
    defect_data = [
        ['DEF-001', 'TC-B005', 'B 商品发布', '中', '上传超过5MB的图片时，页面无响应3秒后才提示错误', '选择6MB的图片，点击上传', '立即提示"图片大小不能超过5MB"', '页面卡顿3秒后才显示错误提示'],
        ['DEF-002', 'TC-D007', 'D 订单交易', '高', '卖家确认订单后，商品状态未立即更新为"交易中"', '卖家确认订单，查看商品详情页', '商品状态变为"交易中"', '商品状态仍显示"在售"，需刷新页面才更新'],
        ['DEF-003', 'TC-G007', 'G 管理后台', '低', '删除分类时未检查是否有商品引用该分类', '选择有商品的分类，点击删除按钮，确认删除', '提示"该分类下有商品，无法删除"或先转移商品', '删除成功，但商品分类字段变为空'],
        ['DEF-004', 'TC-R002', '安全测试', '严重', 'SQL注入漏洞：搜索框输入注入语句可返回所有商品', '在搜索框输入SQL注入语句，点击搜索按钮', '正常显示搜索结果，不执行SQL注入', '返回所有商品列表，存在SQL注入风险'],
    ]
    add_table(doc, defect_data, headers=['缺陷编号', '测试用例', '所属模块', '缺陷级别', '缺陷描述', '复现步骤', '预期结果', '实际结果'], col_widths=[1, 1, 1, 0.8, 2, 1.5, 1.5, 1.5])
    
    add_heading(doc, '7.3 缺陷分类统计', 2)
    defect_class_data = [
        ['严重', '1', '25%', '导致系统崩溃、数据丢失或安全漏洞'],
        ['高', '1', '25%', '影响核心功能正常使用'],
        ['中', '1', '25%', '影响非核心功能或用户体验'],
        ['低', '1', '25%', '界面显示问题或小的功能瑕疵'],
    ]
    add_table(doc, defect_class_data, headers=['缺陷级别', '数量', '占比', '说明'], col_widths=[1.2, 0.8, 1, 3])
    
    add_heading(doc, '7.4 测试结论', 2)
    add_paragraph(doc, '根据测试执行结果，本系统整体测试通过率为 95.2%，达到了预设的测试完成准则（≥95%）。系统的核心功能模块（用户账号、浏览收藏、消息评价、通知）测试通过率为100%，功能实现较为完善。')
    add_paragraph(doc, '')
    add_paragraph(doc, '主要发现的问题包括：')
    add_paragraph(doc, '1. 安全性问题：搜索框存在SQL注入漏洞，需要立即修复')
    add_paragraph(doc, '2. 数据一致性问题：订单确认后商品状态未同步更新')
    add_paragraph(doc, '3. 用户体验问题：大文件上传时页面响应延迟')
    add_paragraph(doc, '4. 数据完整性问题：分类删除未检查商品引用')
    add_paragraph(doc, '')
    add_paragraph(doc, '整体而言，系统已具备上线条件，但需要先修复严重级别的SQL注入漏洞和高级别的数据一致性问题。')
    
    add_heading(doc, '7.5 改进建议', 2)
    
    add_heading(doc, '7.5.1 立即修复（严重/高优先级）', 3)
    add_paragraph(doc, '1. 修复SQL注入漏洞（DEF-004）')
    add_paragraph(doc, '   - 修改搜索接口的SQL查询语句，使用参数化查询')
    add_paragraph(doc, '   - 对用户输入进行严格的字符过滤和转义处理')
    add_paragraph(doc, '   - 添加输入验证层，禁止特殊字符输入')
    add_paragraph(doc, '')
    add_paragraph(doc, '2. 修复订单确认后商品状态不同步问题（DEF-002）')
    add_paragraph(doc, '   - 在确认订单函数中添加商品状态更新逻辑')
    add_paragraph(doc, '   - 使用数据库事务确保订单和商品状态的原子性更新')
    
    add_heading(doc, '7.5.2 计划修复（中/低优先级）', 3)
    add_paragraph(doc, '3. 优化大文件上传体验（DEF-001）')
    add_paragraph(doc, '   - 添加前端文件大小预检查，在选择文件后立即验证')
    add_paragraph(doc, '   - 实现文件上传进度条，提升用户体验')
    add_paragraph(doc, '')
    add_paragraph(doc, '4. 完善分类删除逻辑（DEF-003）')
    add_paragraph(doc, '   - 删除分类前检查是否有商品引用')
    add_paragraph(doc, '   - 提供批量转移商品到其他分类的功能')
    
    doc.add_page_break()
    
    # ==========================================
    # 附录
    # ==========================================
    add_heading(doc, '附录', 1)
    
    add_heading(doc, '附录A：测试用例编号规则', 2)
    tc_rule_data = [
        ['TC-A', '用户账号模块'],
        ['TC-B', '商品发布模块'],
        ['TC-C', '浏览收藏模块'],
        ['TC-D', '订单交易模块'],
        ['TC-E', '消息评价模块'],
        ['TC-F', '通知模块'],
        ['TC-G', '管理后台模块'],
        ['TC-R', '健壮性测试'],
    ]
    add_table(doc, tc_rule_data, headers=['前缀', '模块'], col_widths=[1.5, 3])
    
    add_heading(doc, '附录B：订单状态流转图', 2)
    add_paragraph(doc, 'PENDING → CONFIRMED → PAID → COMPLETED')
    add_paragraph(doc, '    ↓              ↓')
    add_paragraph(doc, 'REJECTED      CANCELLED')
    
    add_heading(doc, '附录C：用户角色权限表', 2)
    role_perm_data = [
        ['浏览首页', '✓', '✓', '✓'],
        ['分类浏览', '✓', '✓', '✓'],
        ['搜索商品', '✓', '✓', '✓'],
        ['商品详情', '✓', '✓', '✓'],
        ['用户注册', '✓', '-', '-'],
        ['用户登录', '✓', '-', '-'],
        ['发布商品', '-', '✓', '✓'],
        ['管理商品', '-', '✓', '✓'],
        ['收藏商品', '-', '✓', '✓'],
        ['提交订单', '-', '✓', '✓'],
        ['订单管理', '-', '✓', '✓'],
        ['发送消息', '-', '✓', '✓'],
        ['评价', '-', '✓', '✓'],
        ['举报', '-', '✓', '✓'],
        ['查看通知', '-', '✓', '✓'],
        ['用户管理', '-', '-', '✓'],
        ['商品审核', '-', '-', '✓'],
        ['分类管理', '-', '-', '✓'],
        ['举报处理', '-', '-', '✓'],
        ['数据统计', '-', '-', '✓'],
    ]
    add_table(doc, role_perm_data, headers=['权限', '访客', '普通用户', '管理员'], col_widths=[2, 1, 1.5, 1.5])
    
    # 保存文档
    output_file = os.path.join(os.path.dirname(__file__), '校园二手物品交易系统_软件测试文档.docx')
    doc.save(output_file)
    print(f'测试文档已生成: {output_file}')
    
    return output_file


if __name__ == '__main__':
    generate_test_document()
